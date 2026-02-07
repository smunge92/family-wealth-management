"""
Generate a professionally formatted Word document for the Architecture Design Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("_" * 80)
    run.font.color.rgb = RGBColor(0, 102, 204)

def create_styled_table(doc, headers, rows, header_color="1F4E79", alt_row_color="D6EAF8"):
    """Create a professionally styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[col_idx], alt_row_color)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table

def add_code_block(doc, code_text):
    """Add a formatted code block"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)

def main():
    doc = Document()

    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = "Family Wealth Management Team"
    core_properties.title = "Architecture Design Document"

    # =========================================================================
    # TITLE PAGE
    # =========================================================================

    # Add some space at top
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FAMILY WEALTH MANAGEMENT")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(31, 78, 121)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Architecture Design Document")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(68, 114, 196)

    # Decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("_" * 50)
    run.font.color.rgb = RGBColor(31, 78, 121)

    # Version info
    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Version 1.0\nFebruary 2026\n\nCONFIDENTIAL")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Page break
    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================

    toc_title = doc.add_paragraph()
    run = toc_title.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    toc_items = [
        ("1.", "Executive Summary", 3),
        ("2.", "System Overview", 4),
        ("3.", "Architecture Diagram", 5),
        ("4.", "Frontend Architecture", 6),
        ("5.", "Backend Architecture", 8),
        ("6.", "Database Design", 11),
        ("7.", "Authentication & Authorization", 14),
        ("8.", "External Integrations", 16),
        ("9.", "Data Flows", 18),
        ("10.", "Security Architecture", 20),
        ("11.", "Deployment Architecture", 22),
        ("12.", "Configuration Management", 23),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run1 = p.add_run(f"{num} ")
        run1.font.bold = True
        run1.font.size = Pt(12)
        run1.font.color.rgb = RGBColor(31, 78, 121)

        run2 = p.add_run(title)
        run2.font.size = Pt(12)

        # Add dots
        run3 = p.add_run(" " + "." * (60 - len(title)))
        run3.font.size = Pt(12)
        run3.font.color.rgb = RGBColor(180, 180, 180)

        run4 = p.add_run(f" {page}")
        run4.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("1. EXECUTIVE SUMMARY")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    p.add_run("Family Wealth Management").bold = True
    p.add_run(" is a personal finance application designed for family-only access. It enables users to:")

    features = [
        "Connect bank accounts via Plaid for automated transaction sync",
        "Import historical data from CSV/Excel/OFX files",
        "Track net worth across all accounts with daily snapshots",
        "Analyze spending patterns with category breakdowns",
        "Plan for retirement using Monte Carlo simulations",
        "Receive AI-powered insights from Claude for financial decisions"
    ]

    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature)

    doc.add_paragraph()

    # Technology Stack Table
    h2 = doc.add_paragraph()
    run = h2.add_run("Technology Stack Summary")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    tech_stack = [
        ("Frontend", "React 18, TypeScript, MSAL, Recharts"),
        ("Backend", "Azure Functions (Python 3.11+)"),
        ("Database", "Azure SQL Server"),
        ("Authentication", "Azure Entra ID (Azure AD)"),
        ("Bank Data", "Plaid API"),
        ("AI Analysis", "Anthropic Claude API"),
        ("Infrastructure", "Azure (Functions, SQL, Key Vault, Storage)"),
    ]

    create_styled_table(doc, ["Layer", "Technology"], tech_stack)

    doc.add_page_break()

    # =========================================================================
    # SECTION 2: SYSTEM OVERVIEW
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("2. SYSTEM OVERVIEW")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    p.add_run("The system follows a modern three-tier architecture with a React single-page application (SPA) frontend, serverless Azure Functions backend, and Azure SQL database.")

    doc.add_paragraph()

    # High-level architecture diagram as ASCII
    h2 = doc.add_paragraph()
    run = h2.add_run("High-Level Architecture")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    architecture_ascii = """
+-------------------+     +-------------------+     +-------------------+
|                   |     |                   |     |                   |
|   USER BROWSER    | --> |  AZURE FUNCTIONS  | --> |    AZURE SQL      |
|   (React SPA)     |     |    (Python)       |     |   (Database)      |
|                   |     |                   |     |                   |
+-------------------+     +-------------------+     +-------------------+
                                   |
                    +--------------+--------------+
                    |              |              |
                    v              v              v
            +------------+  +------------+  +------------+
            | Plaid API  |  | Claude API |  | Azure AD   |
            | (Banking)  |  |   (AI)     |  |  (Auth)    |
            +------------+  +------------+  +------------+
"""
    add_code_block(doc, architecture_ascii)

    doc.add_paragraph()

    # Key Components
    h2 = doc.add_paragraph()
    run = h2.add_run("Key Components")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    components = [
        ("React SPA", "Single-page application with TypeScript, handling user interface and client-side state management"),
        ("Azure Functions", "Serverless Python functions providing RESTful API endpoints"),
        ("Azure SQL", "Relational database storing users, accounts, transactions, and balances"),
        ("Plaid API", "Third-party service for secure bank account connection and transaction sync"),
        ("Claude API", "Anthropic's AI for generating financial insights and analysis"),
        ("Azure Entra ID", "Identity provider for OAuth 2.0 authentication"),
    ]

    create_styled_table(doc, ["Component", "Description"], components)

    doc.add_page_break()

    # =========================================================================
    # SECTION 3: ARCHITECTURE DIAGRAM
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("3. ARCHITECTURE DIAGRAM")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("Layered Architecture")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    layers = [
        ("Presentation Layer", "React components (Dashboard, Accounts, Transactions, Insights, Planning)"),
        ("Service Layer", "MSAL authentication, Axios API client, React Context state"),
        ("API Gateway", "Azure Functions HTTP triggers with routing"),
        ("Business Logic", "Plaid integration, Claude analysis, Data aggregation"),
        ("Data Access", "Database manager, Encryption, Validation, Rate limiting"),
        ("Data Storage", "Azure SQL (users, accounts, transactions, balances, insights)"),
    ]

    create_styled_table(doc, ["Layer", "Components"], layers)

    doc.add_paragraph()

    # Component flow
    h2 = doc.add_paragraph()
    run = h2.add_run("Request Flow")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    flow = """
User Action --> React Component --> API Service (Axios)
                                          |
                                    Bearer Token Added
                                          |
                                          v
                            Azure Functions (@require_auth)
                                          |
                            +-------------+-------------+
                            |             |             |
                            v             v             v
                      Validation    Rate Limit    User Isolation
                            |             |             |
                            +-------------+-------------+
                                          |
                                          v
                                 Business Logic Layer
                                          |
                            +-------------+-------------+
                            |             |             |
                            v             v             v
                       Plaid API    Claude API    Database
"""
    add_code_block(doc, flow)

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: FRONTEND ARCHITECTURE
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("4. FRONTEND ARCHITECTURE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("4.1 Technology Stack")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    frontend_tech = [
        ("Framework", "React", "18.2.0"),
        ("Language", "TypeScript", "4.9.5"),
        ("Routing", "React Router DOM", "6.21.1"),
        ("State", "React Context API", "-"),
        ("Auth", "MSAL React", "2.0.11"),
        ("HTTP", "Axios", "1.6.5"),
        ("Charts", "Recharts", "2.15.4"),
        ("Plaid", "React Plaid Link", "3.5.1"),
    ]

    create_styled_table(doc, ["Component", "Technology", "Version"], frontend_tech)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("4.2 Directory Structure")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    dir_structure = """
frontend/src/
+-- components/
|   +-- Dashboard/       # Net worth & analytics dashboard
|   +-- Accounts/        # Account management & Plaid connection
|   +-- Transactions/    # Transaction history & analysis
|   +-- Insights/        # AI-powered financial insights
|   +-- Planning/        # Retirement & house planning
|   +-- Auth/            # Login component
|   +-- common/          # Reusable components (Toast, Filters)
+-- services/
|   +-- auth.ts          # MSAL configuration
|   +-- api.ts           # Axios instance with interceptors
+-- context/
|   +-- FamilyMemberContext.tsx  # Family member state
+-- App.tsx              # Main app with routing
+-- index.tsx            # Entry point with MSAL
"""
    add_code_block(doc, dir_structure)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("4.3 Page Components")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    pages = [
        ("Dashboard", "Net worth overview, recent transactions, balance charts, spending trends"),
        ("Accounts", "Connected accounts list, sync status, Plaid connection modal"),
        ("Transactions", "Paginated history, category filtering, date range, CSV export"),
        ("Insights", "AI portfolio analysis, spending analytics, retirement/house planning"),
        ("Planning", "Monte Carlo retirement simulations, house affordability calculator"),
    ]

    create_styled_table(doc, ["Page", "Features"], pages)

    doc.add_page_break()

    h2 = doc.add_paragraph()
    run = h2.add_run("4.4 State Management")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    p = doc.add_paragraph()
    p.add_run("The application uses React Context API for global state management. The primary context is ")
    p.add_run("FamilyMemberContext").bold = True
    p.add_run(" which provides:")

    state_items = [
        "familyMembers: Array of family members for the current user",
        "selectedMemberId: Currently selected family member (null = all family)",
        "setSelectedMemberId(): Update selection (persisted to localStorage)",
        "addFamilyMember(): POST to /family-members API",
        "deleteFamilyMember(): DELETE family member",
    ]

    for item in state_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: BACKEND ARCHITECTURE
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("5. BACKEND ARCHITECTURE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("5.1 Technology Stack")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    backend_tech = [
        ("Runtime", "Azure Functions v2"),
        ("Language", "Python 3.11+"),
        ("Database", "pymssql / SQLAlchemy"),
        ("Auth", "MSAL 1.24+"),
        ("Encryption", "cryptography (AES-256-GCM)"),
    ]

    create_styled_table(doc, ["Component", "Technology"], backend_tech)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("5.2 Function Modules")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    modules = [
        ("plaid_integration/", "Bank connection, transaction sync, webhooks, CSV import"),
        ("claude_integration/", "AI portfolio analysis, retirement planning, spending insights"),
        ("data_aggregation/", "Balance aggregation, net worth calculation, tax integration"),
        ("api/", "CRUD endpoints for accounts, transactions, family members"),
        ("scheduled_jobs/", "Daily sync timer, weekly analysis timer"),
    ]

    create_styled_table(doc, ["Module", "Purpose"], modules)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("5.3 API Endpoints")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    endpoints = [
        ("/plaid/link-token", "POST", "Create Plaid Link token", "5/min"),
        ("/plaid/exchange-token", "POST", "Exchange public token", "-"),
        ("/plaid/webhook", "POST", "Plaid webhooks", "Signature"),
        ("/transactions/sync", "POST", "Sync transactions", "2/min"),
        ("/accounts", "GET", "List accounts", "60/min"),
        ("/accounts/{id}", "DELETE", "Delete account", "10/min"),
        ("/transactions", "GET", "Query transactions", "60/min"),
        ("/family-members", "GET/POST", "List/Create members", "60/min"),
        ("/insights/portfolio", "POST", "AI portfolio analysis", "10/min"),
        ("/insights/retirement", "POST", "Retirement planning", "10/min"),
    ]

    create_styled_table(doc, ["Endpoint", "Method", "Function", "Rate Limit"], endpoints)

    doc.add_page_break()

    h2 = doc.add_paragraph()
    run = h2.add_run("5.4 Shared Modules")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    shared = [
        ("auth.py", "JWT validation, user isolation, CORS headers, @require_auth decorator"),
        ("database.py", "SQL connection manager, CRUD operations, query execution"),
        ("encryption.py", "AES-256-GCM encrypt/decrypt for Plaid access tokens"),
        ("validation.py", "Input sanitization, type validation, security checks"),
        ("rate_limiter.py", "In-memory/Redis request throttling"),
        ("audit.py", "Security event logging for compliance"),
        ("plaid_client.py", "Plaid API wrapper (link, exchange, sync)"),
        ("claude_client.py", "Anthropic Claude API wrapper"),
        ("monte_carlo.py", "Retirement projection simulations"),
    ]

    create_styled_table(doc, ["Module", "Purpose"], shared)

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: DATABASE DESIGN
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("6. DATABASE DESIGN")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("6.1 Entity Relationship Overview")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    er_diagram = """
+--------+       +--------------+       +-------------+
| users  |<------| family_      |       | institutions|
+--------+       | members      |       +-------------+
    |            +--------------+              |
    |                   |                      |
    v                   v                      v
+----------------------------------------------------------+
|                       accounts                            |
| (user_id, family_member_id, institution_id, plaid_token) |
+----------------------------------------------------------+
           |                              |
           v                              v
   +---------------+              +---------------+
   | transactions  |              |   balances    |
   | (account_id)  |              | (account_id)  |
   +---------------+              +---------------+
"""
    add_code_block(doc, er_diagram)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("6.2 Core Tables")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    tables = [
        ("users", "user_id (PK)", "Azure Entra ID user OID, email, name"),
        ("family_members", "family_member_id (PK)", "Family members linked to a user"),
        ("institutions", "institution_id (PK)", "Bank/financial institution metadata"),
        ("accounts", "account_id (PK)", "Connected accounts with encrypted tokens"),
        ("transactions", "transaction_id (PK)", "All transactions (Plaid + CSV)"),
        ("balances", "balance_id (PK)", "Daily balance snapshots per account"),
        ("ai_insights", "insight_id (PK)", "Cached AI analysis responses"),
        ("data_imports", "import_id (PK)", "CSV/Excel import tracking"),
        ("plaid_sync_cursors", "cursor_id (PK)", "Plaid sync state per account"),
    ]

    create_styled_table(doc, ["Table", "Primary Key", "Description"], tables)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("6.3 Key Relationships")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    relationships = [
        ("users -> family_members", "One-to-Many", "A user has multiple family members"),
        ("users -> accounts", "One-to-Many", "A user has multiple accounts"),
        ("family_members -> accounts", "One-to-Many", "Family member owns accounts"),
        ("accounts -> transactions", "One-to-Many (CASCADE)", "Account has transactions"),
        ("accounts -> balances", "One-to-Many (CASCADE)", "Daily balance snapshots"),
    ]

    create_styled_table(doc, ["Relationship", "Type", "Description"], relationships)

    doc.add_page_break()

    h2 = doc.add_paragraph()
    run = h2.add_run("6.4 Database Views")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    views = [
        ("vw_current_balances", "Latest balance per account with user info"),
        ("vw_net_worth", "User net worth aggregation"),
        ("vw_monthly_spending", "Spending by category per month"),
        ("vw_account_summary", "Account details with transaction counts"),
        ("vw_transaction_history", "Transactions with account/institution context"),
        ("vw_net_worth_history", "Daily net worth time-series"),
        ("vw_monthly_cashflow", "Income vs expenses by month"),
    ]

    create_styled_table(doc, ["View", "Purpose"], views)

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: AUTHENTICATION & AUTHORIZATION
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("7. AUTHENTICATION & AUTHORIZATION")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("7.1 Authentication Flow")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    auth_flow = """
1. User clicks "Sign In" in React app
2. MSAL redirects to Azure Entra ID login
3. User authenticates with Microsoft
4. Azure returns tokens (access, id, refresh)
5. MSAL stores tokens in sessionStorage
6. React app renders authenticated content
7. API requests include Bearer token
8. Backend validates JWT with Azure JWKS
9. Backend checks email in ALLOWED_USERS
10. Backend validates user can access resource
"""
    add_code_block(doc, auth_flow)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("7.2 Authorization Layers")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    auth_layers = [
        ("Token Validation", "JWT signature, expiry, audience, issuer", "401 Unauthorized"),
        ("User Allowlist", "Email in ALLOWED_USERS env var", "403 Forbidden"),
        ("User Isolation", "Token user_id matches requested user_id", "403 Forbidden"),
    ]

    create_styled_table(doc, ["Layer", "Check", "Failure Response"], auth_layers)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("7.3 Security Headers")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    headers = [
        ("Strict-Transport-Security", "max-age=31536000; includeSubDomains"),
        ("X-Frame-Options", "DENY"),
        ("X-Content-Type-Options", "nosniff"),
        ("X-XSS-Protection", "1; mode=block"),
        ("Content-Security-Policy", "default-src 'self'; frame-ancestors 'none'"),
        ("Cache-Control", "no-store, no-cache, must-revalidate"),
    ]

    create_styled_table(doc, ["Header", "Value"], headers)

    doc.add_page_break()

    # =========================================================================
    # SECTION 8: EXTERNAL INTEGRATIONS
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("8. EXTERNAL INTEGRATIONS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("8.1 Plaid API")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    p = doc.add_paragraph()
    p.add_run("Purpose: ").bold = True
    p.add_run("Secure bank account connection and transaction synchronization")

    plaid_endpoints = [
        ("/link/token/create", "Create link token for frontend modal"),
        ("/item/public_token/exchange", "Exchange public token for access token"),
        ("/transactions/sync", "Incremental transaction sync (cursor-based)"),
        ("/accounts/get", "List accounts for an item"),
        ("/institutions/get_by_id", "Get institution metadata"),
    ]

    create_styled_table(doc, ["Plaid Endpoint", "Purpose"], plaid_endpoints)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("8.2 Anthropic Claude API")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    p = doc.add_paragraph()
    p.add_run("Purpose: ").bold = True
    p.add_run("AI-powered financial analysis and personalized insights")

    claude_uses = [
        ("Portfolio Analysis", "Asset allocation, risk profile, diversification recommendations"),
        ("Retirement Planning", "Monte Carlo projections, savings rate optimization"),
        ("House Affordability", "Purchase feasibility, timeline, down payment analysis"),
        ("Spending Analysis", "Category breakdown, optimization opportunities"),
    ]

    create_styled_table(doc, ["Analysis Type", "Description"], claude_uses)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("8.3 Azure Services")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    azure_services = [
        ("Azure Entra ID", "User authentication (OAuth 2.0 / OIDC)"),
        ("Azure SQL Database", "Primary relational data store"),
        ("Azure Key Vault", "Secrets management (API keys, credentials)"),
        ("Azure Functions", "Serverless API hosting"),
        ("Application Insights", "Monitoring and logging"),
    ]

    create_styled_table(doc, ["Service", "Purpose"], azure_services)

    doc.add_page_break()

    # =========================================================================
    # SECTION 9: DATA FLOWS
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("9. DATA FLOWS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("9.1 Account Connection Flow")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    account_flow = """
1. User clicks "Connect Account" in React app
2. Frontend requests link token from backend
3. Backend calls Plaid /link/token/create
4. Frontend displays Plaid Link modal
5. User authenticates with their bank
6. Plaid returns public token to frontend
7. Frontend sends public token to backend
8. Backend exchanges for access token (encrypted)
9. Backend fetches accounts and balances
10. Backend stores in database
11. Backend returns account list to frontend
"""
    add_code_block(doc, account_flow)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("9.2 Transaction Sync Flow")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    sync_flow = """
1. Timer trigger (daily) or user request
2. For each active account:
   a. Get cursor from plaid_sync_cursors
   b. Call Plaid /transactions/sync
   c. Process added/modified/removed transactions
   d. Update cursor value
   e. Aggregate daily balance
3. Update net worth views
"""
    add_code_block(doc, sync_flow)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("9.3 AI Analysis Flow")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    ai_flow = """
1. User requests portfolio analysis
2. Backend checks cache (24h TTL)
3. If cache miss:
   a. Fetch accounts, transactions, balances
   b. Build Claude prompt with financial data
   c. Call Claude API
   d. Store response in ai_insights
4. Return analysis to frontend
"""
    add_code_block(doc, ai_flow)

    doc.add_page_break()

    # =========================================================================
    # SECTION 10: SECURITY ARCHITECTURE
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("10. SECURITY ARCHITECTURE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("10.1 Security Layers")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    security_layers = [
        ("Layer 1: Transport", "TLS 1.2+ for all connections, HSTS header enforces HTTPS"),
        ("Layer 2: Authentication", "Azure Entra ID (OAuth 2.0), JWT validation with JWKS"),
        ("Layer 3: Authorization", "Email allowlist, user isolation, rate limiting"),
        ("Layer 4: Data Protection", "AES-256-GCM encryption, parameterized SQL queries"),
        ("Layer 5: Audit", "All requests logged, security events tracked, no secrets in logs"),
    ]

    create_styled_table(doc, ["Layer", "Implementation"], security_layers)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("10.2 Encryption")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    encryption = [
        ("Plaid Access Tokens", "At rest", "AES-256-GCM (Fernet)"),
        ("Database Connection", "In transit", "TLS 1.2+"),
        ("API Requests", "In transit", "HTTPS"),
    ]

    create_styled_table(doc, ["Data", "Protection", "Algorithm"], encryption)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("10.3 Rate Limiting")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    rate_limits = [
        ("Read operations", "60 requests", "1 minute"),
        ("Write operations", "30 requests", "1 minute"),
        ("Delete operations", "10 requests", "1 minute"),
        ("Plaid link", "5 requests", "1 minute"),
        ("AI insights", "10 requests", "1 minute"),
    ]

    create_styled_table(doc, ["Category", "Limit", "Window"], rate_limits)

    doc.add_page_break()

    # =========================================================================
    # SECTION 11: DEPLOYMENT ARCHITECTURE
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("11. DEPLOYMENT ARCHITECTURE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("11.1 Azure Resources")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    azure_resources = [
        ("Azure Function App", "Python 3.11 runtime, serverless API hosting"),
        ("Azure SQL Database", "Relational data store, automated backups"),
        ("Azure Key Vault", "Secrets management, encryption keys"),
        ("Azure Entra ID", "Identity provider, app registrations"),
        ("Application Insights", "Monitoring, logging, diagnostics"),
        ("Azure Storage", "Frontend static hosting (optional)"),
    ]

    create_styled_table(doc, ["Resource", "Purpose"], azure_resources)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("11.2 Environment Configuration")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    environments = [
        ("Local Dev", "localhost:3000", "localhost:7071", "Azure SQL (remote)"),
        ("Production", "https://app.domain.com", "https://api.domain.com", "Azure SQL"),
    ]

    create_styled_table(doc, ["Environment", "Frontend URL", "Backend URL", "Database"], environments)

    doc.add_page_break()

    # =========================================================================
    # SECTION 12: CONFIGURATION MANAGEMENT
    # =========================================================================

    h1 = doc.add_paragraph()
    run = h1.add_run("12. CONFIGURATION MANAGEMENT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_horizontal_line(doc)

    h2 = doc.add_paragraph()
    run = h2.add_run("12.1 Frontend Environment Variables")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    frontend_env = [
        ("REACT_APP_AZURE_CLIENT_ID", "Azure app registration client ID"),
        ("REACT_APP_AZURE_TENANT_ID", "Azure tenant ID"),
        ("REACT_APP_AZURE_REDIRECT_URI", "OAuth redirect URI"),
        ("REACT_APP_API_BASE_URL", "Backend API base URL"),
        ("REACT_APP_PLAID_ENV", "Plaid environment (sandbox/production)"),
    ]

    create_styled_table(doc, ["Variable", "Purpose"], frontend_env)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run("12.2 Backend Environment Variables")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    backend_env = [
        ("AZURE_TENANT_ID", "Azure tenant ID"),
        ("AZURE_CLIENT_ID", "Azure client ID"),
        ("AZURE_CLIENT_SECRET", "Azure client secret (from Key Vault)"),
        ("SQL_SERVER", "Azure SQL server hostname"),
        ("SQL_DATABASE", "Database name"),
        ("PLAID_CLIENT_ID", "Plaid API client ID"),
        ("PLAID_SECRET", "Plaid API secret (from Key Vault)"),
        ("ANTHROPIC_API_KEY", "Claude API key (from Key Vault)"),
        ("REQUIRE_AUTH", "Enable authentication (true/false)"),
        ("ALLOWED_USERS", "Comma-separated email allowlist"),
        ("ENCRYPTION_KEY", "Encryption key (from Key Vault)"),
    ]

    create_styled_table(doc, ["Variable", "Purpose"], backend_env)

    # =========================================================================
    # FOOTER
    # =========================================================================

    doc.add_page_break()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("Family Wealth Management - Architecture Design Document\nVersion 1.0 | February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), "Architecture_Design_Document.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
